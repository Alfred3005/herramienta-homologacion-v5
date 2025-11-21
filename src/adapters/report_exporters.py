"""
Exportadores de Reportes a Múltiples Formatos
Soporta: PDF, DOCX, HTML, TXT

Versión: 1.0
Fecha: 2025-11-20
"""

from typing import Dict, Any, Optional
from pathlib import Path
from datetime import datetime
import io


class ReportExporterBase:
    """Clase base para exportadores de reportes"""

    def __init__(self):
        self.formato = "base"

    def exportar(self, contenido_reporte: str, metadata: Optional[Dict[str, Any]] = None) -> bytes:
        """
        Exporta reporte a formato específico.

        Args:
            contenido_reporte: Texto del reporte
            metadata: Metadatos opcionales (título, código puesto, etc.)

        Returns:
            Bytes del archivo generado
        """
        raise NotImplementedError("Subclases deben implementar exportar()")


class TXTExporter(ReportExporterBase):
    """Exportador a formato TXT plano"""

    def __init__(self):
        super().__init__()
        self.formato = "txt"

    def exportar(self, contenido_reporte: str, metadata: Optional[Dict[str, Any]] = None) -> bytes:
        """Exporta a TXT con encoding UTF-8"""
        # Agregar header si hay metadata
        if metadata:
            header = f"""{'='*70}
REPORTE DE PUESTO - FORMATO RH NET
Código: {metadata.get('codigo_puesto', 'N/A')}
Generado: {metadata.get('fecha_generacion', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))}
{'='*70}

"""
            contenido_completo = header + contenido_reporte
        else:
            contenido_completo = contenido_reporte

        return contenido_completo.encode('utf-8')


class HTMLExporter(ReportExporterBase):
    """Exportador a formato HTML con estilos"""

    def __init__(self):
        super().__init__()
        self.formato = "html"

    def exportar(self, contenido_reporte: str, metadata: Optional[Dict[str, Any]] = None) -> bytes:
        """Exporta a HTML con estilos CSS"""
        # Convertir saltos de línea a <br>
        contenido_html = contenido_reporte.replace('\n', '<br>\n')

        # Template HTML
        html = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Reporte RH Net - {metadata.get('codigo_puesto', 'Puesto') if metadata else 'Puesto'}</title>
    <style>
        body {{
            font-family: 'Arial', sans-serif;
            max-width: 900px;
            margin: 40px auto;
            padding: 20px;
            background-color: #f5f5f5;
            line-height: 1.6;
        }}
        .container {{
            background-color: white;
            padding: 30px;
            border-radius: 8px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }}
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px;
            border-radius: 8px 8px 0 0;
            margin: -30px -30px 20px -30px;
        }}
        .header h1 {{
            margin: 0;
            font-size: 24px;
        }}
        .header p {{
            margin: 5px 0 0 0;
            opacity: 0.9;
            font-size: 14px;
        }}
        .content {{
            font-size: 14px;
            color: #333;
        }}
        .footer {{
            margin-top: 30px;
            padding-top: 20px;
            border-top: 2px solid #eee;
            text-align: center;
            font-size: 12px;
            color: #888;
        }}
        strong {{
            color: #667eea;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🏛️ Reporte de Puesto - Formato RH Net</h1>
            <p><strong>Código:</strong> {metadata.get('codigo_puesto', 'N/A') if metadata else 'N/A'} |
               <strong>Generado:</strong> {metadata.get('fecha_generacion', datetime.now().strftime('%Y-%m-%d %H:%M:%S')) if metadata else datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </div>
        <div class="content">
            {contenido_html}
        </div>
        <div class="footer">
            Sistema de Homologación APF v5.41 | Generado con Claude Code
        </div>
    </div>
</body>
</html>"""

        return html.encode('utf-8')


class PDFExporter(ReportExporterBase):
    """Exportador a formato PDF usando fpdf2"""

    def __init__(self):
        super().__init__()
        self.formato = "pdf"

    def exportar(self, contenido_reporte: str, metadata: Optional[Dict[str, Any]] = None) -> bytes:
        """Exporta a PDF usando fpdf2"""
        try:
            from fpdf import FPDF
        except ImportError:
            raise ImportError(
                "fpdf2 no está instalado. Instala con: pip install fpdf2"
            )

        # Crear PDF con márgenes adecuados
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_margins(left=15, top=15, right=15)

        # Header con color
        pdf.set_fill_color(102, 126, 234)  # Color morado
        pdf.set_text_color(255, 255, 255)  # Blanco
        pdf.set_font('Arial', 'B', 14)
        pdf.multi_cell(0, 10, 'Reporte de Puesto - Formato RH Net', 0, 'C', True)

        # Metadata
        if metadata:
            pdf.set_fill_color(240, 240, 240)  # Gris claro
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('Arial', '', 9)
            codigo = metadata.get('codigo_puesto', 'N/A')
            fecha = metadata.get('fecha_generacion', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            # Usar multi_cell para evitar problemas de ancho
            pdf.multi_cell(0, 6, f'Código: {codigo} | Generado: {fecha}', 0, 'C', True)

        pdf.ln(5)

        # Contenido
        pdf.set_text_color(0, 0, 0)
        pdf.set_font('Arial', '', 9)

        # Procesar líneas
        for linea in contenido_reporte.split('\n'):
            if not linea.strip():
                pdf.ln(2)
                continue

            # Limitar longitud de línea para evitar problemas
            linea_limpia = linea.strip()

            # Detectar títulos (líneas cortas en mayúsculas)
            if linea_limpia.isupper() and len(linea_limpia) < 50:
                pdf.set_font('Arial', 'B', 10)
                pdf.ln(2)
                pdf.multi_cell(0, 5, linea_limpia, 0, 'L')
                pdf.set_font('Arial', '', 9)
            elif linea.startswith('Función '):
                pdf.set_font('Arial', 'B', 9)
                pdf.ln(1)
                pdf.multi_cell(0, 5, linea_limpia, 0, 'L')
                pdf.set_font('Arial', '', 9)
            else:
                # Texto normal - siempre usar multi_cell para manejo automático de líneas largas
                try:
                    pdf.multi_cell(0, 4.5, linea_limpia, 0, 'L')
                except Exception as e:
                    # Si falla, intentar con línea truncada
                    pdf.multi_cell(0, 4.5, linea_limpia[:180] + '...', 0, 'L')

        # Footer
        pdf.ln(8)
        pdf.set_font('Arial', 'I', 7)
        pdf.set_text_color(128, 128, 128)
        pdf.multi_cell(0, 4, 'Sistema de Homologación APF v5.42 | Generado con Claude Code', 0, 'C')

        # Retornar bytes
        return pdf.output()


class DOCXExporter(ReportExporterBase):
    """Exportador a formato DOCX usando python-docx"""

    def __init__(self):
        super().__init__()
        self.formato = "docx"

    def exportar(self, contenido_reporte: str, metadata: Optional[Dict[str, Any]] = None) -> bytes:
        """Exporta a DOCX usando python-docx"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError(
                "python-docx no está instalado. Instala con: pip install python-docx"
            )

        # Crear documento
        doc = Document()

        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Header con título
        header = doc.add_heading('🏛️ Reporte de Puesto - Formato RH Net', level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.runs[0]
        header_run.font.color.rgb = RGBColor(102, 126, 234)

        # Metadata
        if metadata:
            meta_para = doc.add_paragraph()
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            codigo = metadata.get('codigo_puesto', 'N/A')
            fecha = metadata.get('fecha_generacion', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            meta_run = meta_para.add_run(f'Código: {codigo} | Generado: {fecha}')
            meta_run.font.size = Pt(10)
            meta_run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()  # Espacio

        # Procesar contenido
        for linea in contenido_reporte.split('\n'):
            if not linea.strip():
                doc.add_paragraph()  # Línea vacía
                continue

            # Detectar secciones principales (mayúsculas)
            if linea.strip().isupper() and len(linea.strip()) < 50:
                heading = doc.add_heading(linea.strip(), level=2)
                heading_run = heading.runs[0]
                heading_run.font.color.rgb = RGBColor(102, 126, 234)
            # Detectar funciones
            elif linea.startswith('Función '):
                para = doc.add_paragraph()
                run = para.add_run(linea.strip())
                run.bold = True
                run.font.size = Pt(11)
            # Texto normal
            else:
                para = doc.add_paragraph(linea.strip())
                para_format = para.paragraph_format
                para_format.line_spacing = 1.15

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run('Sistema de Homologación APF v5.41 | Generado con Claude Code')
        footer_run.font.size = Pt(8)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)

        # Guardar en bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()


class ReportExporterFactory:
    """Factory para crear exportadores según formato"""

    _exporters = {
        'txt': TXTExporter,
        'html': HTMLExporter,
        'pdf': PDFExporter,
        'docx': DOCXExporter
    }

    @classmethod
    def get_exporter(cls, formato: str) -> ReportExporterBase:
        """
        Obtiene exportador para formato especificado.

        Args:
            formato: Formato deseado ('txt', 'html', 'pdf', 'docx')

        Returns:
            Instancia del exportador

        Raises:
            ValueError: Si formato no es soportado
        """
        formato = formato.lower()
        if formato not in cls._exporters:
            raise ValueError(
                f"Formato '{formato}' no soportado. "
                f"Formatos válidos: {', '.join(cls._exporters.keys())}"
            )

        return cls._exporters[formato]()

    @classmethod
    def formatos_disponibles(cls) -> list:
        """Retorna lista de formatos disponibles"""
        return list(cls._exporters.keys())


def exportar_reporte(contenido_reporte: str, formato: str,
                    metadata: Optional[Dict[str, Any]] = None) -> bytes:
    """
    Función helper para exportar reporte a formato especificado.

    Args:
        contenido_reporte: Texto del reporte
        formato: Formato deseado ('txt', 'html', 'pdf', 'docx')
        metadata: Metadatos opcionales

    Returns:
        Bytes del archivo generado

    Example:
        >>> reporte = generador.generar_reporte_completo(datos)
        >>> pdf_bytes = exportar_reporte(reporte, 'pdf', {'codigo_puesto': '12345'})
    """
    exporter = ReportExporterFactory.get_exporter(formato)
    return exporter.exportar(contenido_reporte, metadata)
