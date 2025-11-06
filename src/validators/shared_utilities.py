# -*- coding: utf-8 -*-
"""
SHARED UTILITIES - Sistema APF v4.0 Refactorizado
Utilidades compartidas entre todos los agentes del sistema
Base común unificada con manejo de contexto consolidado
"""

import json
import time
import traceback
import os
import re
import hashlib
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Union, Tuple
from dataclasses import dataclass, field

# Importaciones para procesamiento de documentos
try:
    from litellm import completion
    LITELLM_AVAILABLE = True
except ImportError:
    LITELLM_AVAILABLE = False
    print("Warning: litellm no disponible - funcionalidad OpenAI limitada")

try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("Warning: PyMuPDF no disponible - procesamiento PDF limitado")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx no disponible - procesamiento DOCX limitado")

# ==========================================
# CONFIGURACIÓN GLOBAL DEL SISTEMA
# ==========================================

# Rutas por defecto del sistema
DEFAULT_PATHS = {
    "normativa_dir": "/content/drive/MyDrive/Metodología para la valuación de puestos/Herramienta de homologación/Datos/NORMATIVA-TEXT",
    "schema_file": "/content/drive/MyDrive/Metodología para la valuación de puestos/Herramienta de homologación/Datos/Fomato_extracción/Descripcion_puesto_de_trabajo_todos_los_campos_03.txt",
    "test_puesto": "/content/drive/MyDrive/Metodología para la valuación de puestos/Herramienta de homologación/Datos/Puestos de trabajo/Anticorrupcion/JEFE(A) DE DEPARTAMENTO DE APOYO NORMATIVO EN TRANSPARENCIA.txt"
}

# Configuración de logging
LOGGING_CONFIG = {
    "enable_detailed_logging": True,
    "log_openai_calls": True,
    "log_file_processing": True,
    "log_errors_only": False
}

# ==========================================
# JERARQUÍA DE VERBOS UNIFICADA
# ==========================================

VERB_HIERARCHY = {
    "G": {
        "level_name": "Secretaría de Estado",
        "appropriate_verbs": ["dictar", "normar", "establecer", "representar", "globalizar", "acreditar", 
                             "autorizar", "colaborar", "conducir", "declarar", "detectar", "determinar", 
                             "disponer", "emitir", "evaluar", "fijar", "fungir", "informar", "jerarquizar", 
                             "objetar", "participar", "patrocinar", "proponer", "rendir"],
        "forbidden_verbs": ["ejecutar", "efectuar", "realizar", "tramitar", "operar"],
        "impact_profile": {
            "decision_scope": "strategic_national",
            "budget_range": "strategic", 
            "error_consequences": "systemic",
            "complexity_level": "transformational"
        }
    },
    "H": {
        "level_name": "Subsecretaría/Oficialía Mayor",
        "appropriate_verbs": ["dirigir", "coordinar", "establecer", "normar", "supervisar", "evaluar",
                             "autorizar", "representar", "conducir", "determinar", "emitir", "fijar"],
        "forbidden_verbs": ["ejecutar", "efectuar", "realizar", "tramitar", "operar"],
        "impact_profile": {
            "decision_scope": "strategic_national",
            "budget_range": "strategic",
            "error_consequences": "systemic", 
            "complexity_level": "transformational"
        }
    },
    "J": {
        "level_name": "Jefatura de Unidad",
        "appropriate_verbs": ["administrar", "coordinar", "dirigir", "evaluar", "supervisar", "autorizar",
                             "conducir", "consolidar", "contribuir", "coparticipar", "definir", "designar",
                             "diagnosticar", "dictaminar", "difundir", "divulgar", "establecer", "globalizar"],
        "forbidden_verbs": ["efectuar", "ejecutar", "tramitar", "operar"],
        "impact_profile": {
            "decision_scope": "interinstitutional",
            "budget_range": "major",
            "error_consequences": "strategic",
            "complexity_level": "innovative"
        }
    },
    "K": {
        "level_name": "Dirección General",
        "appropriate_verbs": ["administrar", "coordinar", "dirigir", "evaluar", "supervisar", "autorizar",
                             "conducir", "consolidar", "contribuir", "definir", "diagnosticar", "dictaminar"],
        "forbidden_verbs": ["efectuar", "ejecutar", "tramitar", "operar"],
        "impact_profile": {
            "decision_scope": "interinstitutional",
            "budget_range": "major",
            "error_consequences": "strategic",
            "complexity_level": "strategic"
        }
    },
    "L": {
        "level_name": "Dirección General Adjunta",
        "appropriate_verbs": ["administrar", "coordinar", "dirigir", "evaluar", "supervisar", "contribuir",
                             "aplicar", "asesorar", "autorizar", "conducir", "consolidar"],
        "forbidden_verbs": ["efectuar", "ejecutar", "tramitar"],
        "impact_profile": {
            "decision_scope": "institutional",
            "budget_range": "significant",
            "error_consequences": "strategic",
            "complexity_level": "strategic"
        }
    },
    "M": {
        "level_name": "Dirección de Área/Coordinación",
        "appropriate_verbs": ["coordinar", "supervisar", "elaborar", "implementar", "aplicar", "asesorar",
                             "auxiliar", "colaborar", "compilar", "comprobar", "concentrar", "contribuir",
                             "controlar", "cuantificar", "definir"],
        "forbidden_verbs": ["dictar", "normar", "representar"],
        "impact_profile": {
            "decision_scope": "institutional",
            "budget_range": "significant",
            "error_consequences": "tactical",
            "complexity_level": "analytical"
        }
    },
    "N": {
        "level_name": "Subdirección de Área",
        "appropriate_verbs": ["supervisar", "elaborar", "coordinar", "ejecutar", "analizar", "aplicar",
                             "aportar", "asegurar", "asesorar", "asignar", "actualizar", "administrar"],
        "forbidden_verbs": ["dictar", "normar", "establecer", "representar"],
        "impact_profile": {
            "decision_scope": "institutional",
            "budget_range": "moderate",
            "error_consequences": "tactical",
            "complexity_level": "analytical"
        }
    },
    "O": {
        "level_name": "Jefatura de Departamento",
        "appropriate_verbs": ["ejecutar", "elaborar", "supervisar", "realizar", "analizar", "aplicar",
                             "atender", "asesorar", "asignar", "actualizar", "administrar", "gestionar"],
        "forbidden_verbs": ["dictar", "normar", "establecer", "representar", "globalizar"],
        "impact_profile": {
            "decision_scope": "local",
            "budget_range": "moderate",
            "error_consequences": "operational",
            "complexity_level": "routine"
        }
    },
    "P": {
        "level_name": "Enlace",
        "appropriate_verbs": ["efectuar", "ejecutar", "realizar", "tramitar", "operar", "elaborar",
                             "emplazar", "enlazar", "entablar", "estimar", "estudiar", "expedir",
                             "fincar", "formalizar"],
        "forbidden_verbs": ["dictar", "normar", "establecer", "representar", "globalizar", "autorizar"],
        "impact_profile": {
            "decision_scope": "local",
            "budget_range": "minimal",
            "error_consequences": "operational",
            "complexity_level": "routine"
        }
    }
}

# Verbos débiles prohibidos en todos los niveles
# NOTA: verb_semantic_analyzer puede expandir esta lista con sinónimos detectados por LLM
WEAK_VERBS = [
    # Lista original
    "coadyuvar", "apoyar", "seguir", "resultar", "crear", "ser", "continuar", "auxiliar",
    # Verbos agregados tras calibración baseline 20251023
    "asegurar", "garantizar", "verificar", "gestionar", "hacer", "realizar", "llevar a cabo",
    "mantener", "considerar", "buscar", "procurar", "tratar", "intentar", "lograr"
]

# ==========================================
# MANEJO DE CONTEXTO UNIFICADO
# ==========================================

@dataclass
class ProcessingStep:
    """Representa un paso en el procesamiento"""
    step_name: str
    status: str  # "pending", "running", "completed", "failed"
    start_time: Optional[datetime] = None
    end_time: Optional[datetime] = None
    error: Optional[str] = None
    result_summary: Optional[str] = None

class APFContext:
    """
    Contexto unificado para todo el sistema APF.
    Reemplaza ActionContext y SharedActionContext con funcionalidad consolidada.
    """
    
    def __init__(self, context_id: str = None):
        self.context_id = context_id or self._generate_context_id()
        self.created_at = datetime.now()
        
        # Datos principales
        self.data = {}
        
        # Estado de procesamiento
        self.processing_steps = {}
        
        # Errores y warnings
        self.errors = []
        self.warnings = []
        
        # Metadatos
        self.metadata = {
            "version": "4.0",
            "last_updated": self.created_at,
            "agents_involved": []
        }
    
    def _generate_context_id(self) -> str:
        """Genera ID único para el contexto"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        random_hash = hashlib.md5(str(time.time()).encode()).hexdigest()[:8]
        return f"APF_{timestamp}_{random_hash}"
    
    def set_data(self, key: str, value: Any, agent_name: str = None) -> None:
        """
        Establece un valor en el contexto.
        
        Args:
            key: Clave del dato
            value: Valor a almacenar
            agent_name: Nombre del agente que establece el dato
        """
        self.data[key] = value
        self.metadata["last_updated"] = datetime.now()
        
        if agent_name and agent_name not in self.metadata["agents_involved"]:
            self.metadata["agents_involved"].append(agent_name)
        
        if LOGGING_CONFIG["enable_detailed_logging"]:
            print(f"[CONTEXT {self.context_id}] Set {key}: {type(value).__name__}")
    
    def get_data(self, key: str, default: Any = None) -> Any:
        """
        Obtiene un valor del contexto.
        
        Args:
            key: Clave del dato
            default: Valor por defecto si no existe
            
        Returns:
            Valor almacenado o default
        """
        return self.data.get(key, default)
    
    def start_step(self, step_name: str, agent_name: str = None) -> None:
        """Inicia un paso de procesamiento"""
        step = ProcessingStep(
            step_name=step_name,
            status="running",
            start_time=datetime.now()
        )
        self.processing_steps[step_name] = step
        
        if agent_name and agent_name not in self.metadata["agents_involved"]:
            self.metadata["agents_involved"].append(agent_name)
        
        if LOGGING_CONFIG["enable_detailed_logging"]:
            print(f"[STEP] {step_name} iniciado por {agent_name or 'sistema'}")
    
    def complete_step(self, step_name: str, result_summary: str = None) -> None:
        """Completa un paso de procesamiento"""
        if step_name in self.processing_steps:
            step = self.processing_steps[step_name]
            step.status = "completed"
            step.end_time = datetime.now()
            step.result_summary = result_summary
            
            if LOGGING_CONFIG["enable_detailed_logging"]:
                duration = (step.end_time - step.start_time).total_seconds()
                print(f"[STEP] {step_name} completado en {duration:.2f}s")
    
    def fail_step(self, step_name: str, error: str) -> None:
        """Marca un paso como fallido"""
        if step_name in self.processing_steps:
            step = self.processing_steps[step_name]
            step.status = "failed"
            step.end_time = datetime.now()
            step.error = error
        
        self.add_error(f"Step {step_name} failed: {error}")
    
    def add_error(self, error: str, agent_name: str = None) -> None:
        """Añade un error al contexto"""
        error_entry = {
            "error": error,
            "timestamp": datetime.now(),
            "agent": agent_name
        }
        self.errors.append(error_entry)
        
        if LOGGING_CONFIG["log_errors_only"] or LOGGING_CONFIG["enable_detailed_logging"]:
            print(f"[ERROR] {agent_name or 'Sistema'}: {error}")
    
    def add_warning(self, warning: str, agent_name: str = None) -> None:
        """Añade un warning al contexto"""
        warning_entry = {
            "warning": warning,
            "timestamp": datetime.now(),
            "agent": agent_name
        }
        self.warnings.append(warning_entry)
        
        if LOGGING_CONFIG["enable_detailed_logging"]:
            print(f"[WARNING] {agent_name or 'Sistema'}: {warning}")
    
    def get_summary(self) -> Dict[str, Any]:
        """Obtiene resumen completo del contexto"""
        completed_steps = len([s for s in self.processing_steps.values() if s.status == "completed"])
        failed_steps = len([s for s in self.processing_steps.values() if s.status == "failed"])
        
        return {
            "context_id": self.context_id,
            "created_at": self.created_at.isoformat(),
            "duration": (datetime.now() - self.created_at).total_seconds(),
            "data_keys": list(self.data.keys()),
            "processing_steps": {
                "total": len(self.processing_steps),
                "completed": completed_steps,
                "failed": failed_steps,
                "success_rate": (completed_steps / len(self.processing_steps) * 100) if self.processing_steps else 0
            },
            "errors": len(self.errors),
            "warnings": len(self.warnings),
            "agents_involved": self.metadata["agents_involved"]
        }
    
    def export_full_context(self) -> Dict[str, Any]:
        """Exporta contexto completo para debugging"""
        return {
            "context_id": self.context_id,
            "metadata": self.metadata,
            "data": {k: str(type(v)) if not isinstance(v, (str, int, float, bool, list, dict)) else v 
                    for k, v in self.data.items()},
            "processing_steps": {name: {
                "step_name": step.step_name,
                "status": step.status,
                "start_time": step.start_time.isoformat() if step.start_time else None,
                "end_time": step.end_time.isoformat() if step.end_time else None,
                "error": step.error,
                "result_summary": step.result_summary
            } for name, step in self.processing_steps.items()},
            "errors": self.errors,
            "warnings": self.warnings
        }

# ==========================================
# CLASE BASE PARA AGENTES
# ==========================================

class APFAgent:
    """
    Clase base para todos los agentes del sistema APF.
    Proporciona funcionalidad común y interfaz consistente.
    """
    
    def __init__(self, agent_name: str, context: APFContext = None):
        self.agent_name = agent_name
        self.context = context or APFContext()
        self.version = "4.0"
        self.initialized = False
        
        # Estadísticas del agente
        self.stats = {
            "operations_count": 0,
            "successful_operations": 0,
            "failed_operations": 0,
            "last_operation": None
        }
        
        self.context.metadata["agents_involved"].append(self.agent_name)
        
    def _log(self, message: str, level: str = "INFO") -> None:
        """Log interno del agente"""
        if LOGGING_CONFIG["enable_detailed_logging"] or (level == "ERROR" and LOGGING_CONFIG["log_errors_only"]):
            timestamp = datetime.now().strftime("%H:%M:%S")
            print(f"[{timestamp}] [{self.agent_name}] [{level}] {message}")
    
    def _update_stats(self, success: bool) -> None:
        """Actualiza estadísticas del agente"""
        self.stats["operations_count"] += 1
        if success:
            self.stats["successful_operations"] += 1
        else:
            self.stats["failed_operations"] += 1
        self.stats["last_operation"] = datetime.now().isoformat()
    
    def get_agent_summary(self) -> Dict[str, Any]:
        """Obtiene resumen del agente"""
        return {
            "agent_name": self.agent_name,
            "version": self.version,
            "initialized": self.initialized,
            "stats": self.stats.copy(),
            "context_id": self.context.context_id
        }

# ==========================================
# UTILIDADES DE PROCESAMIENTO DE TEXTO
# ==========================================

def clean_text_for_processing(text: str) -> str:
    """
    Limpia texto para evitar problemas de encoding y parsing JSON.
    Versión optimizada y unificada.
    """
    if not text:
        return ""
    
    # Normalizar encoding
    if isinstance(text, bytes):
        text = text.decode('utf-8', errors='replace')
    
    # Reemplazar caracteres problemáticos
    replacements = {
        '\x00': ' ',  # Null bytes
        '\x1c': ' ',  # File separator
        '\x1d': ' ',  # Group separator  
        '\x1f': ' ',  # Unit separator
        '\x08': '',   # Backspace
        '\x0c': ' ',  # Form feed
        '\r\n': '\n', # Windows line endings
        '\r': '\n'    # Mac line endings
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    # Normalizar espacios en blanco
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    
    return text

def extract_hierarchical_level(codigo_puesto: str = None,
                              nivel_salarial: str = None,
                              denominacion: str = None) -> str:
    """
    Extrae el nivel jerárquico del puesto basado en diversos indicadores.
    VERSIÓN REFACTORIZADA: Usa FlexibleHierarchyExtractor para mayor adaptabilidad

    Mantiene firma compatible con código existente (wrapper)

    Args:
        codigo_puesto: Código del puesto
        nivel_salarial: Nivel salarial
        denominacion: Denominación del puesto

    Returns:
        Nivel jerárquico (G, H, J, K, L, M, N, O, P) o 'UNKNOWN'
    """
    # Lazy import para evitar dependencia circular
    try:
        from flexible_hierarchy_extractor import extract_hierarchical_level_flexible
        return extract_hierarchical_level_flexible(codigo_puesto, nivel_salarial, denominacion)
    except ImportError:
        # Fallback a implementación original si el nuevo módulo no está disponible
        # (para compatibilidad durante migración)

        # Primero intentar con nivel salarial (más confiable)
        if nivel_salarial:
            nivel_upper = nivel_salarial.upper().strip()
            primer_caracter = nivel_upper[0] if nivel_upper else ""

            if primer_caracter in VERB_HIERARCHY:
                return primer_caracter

        # Buscar en código de puesto
        if codigo_puesto:
            codigo_upper = codigo_puesto.upper()
            # Patrón: buscar letra seguida de números en formato típico APF
            patterns = [
                r'-([GHJKLMNOP])\d+[-_]',
                r'[-_]([GHJKLMNOP])\d+[-_]',
                r'^([GHJKLMNOP])\d+',
                r'([GHJKLMNOP])(\d{1,2})[-_]'
            ]

            for pattern in patterns:
                match = re.search(pattern, codigo_upper)
                if match and match.group(1) in VERB_HIERARCHY:
                    return match.group(1)

        # Fallback: análisis de denominación
        if denominacion:
            denominacion_upper = denominacion.upper()

            # Mapeo de palabras clave a niveles jerárquicos
            level_keywords = {
                'G': ['SECRETARI', 'TITULAR'],
                'H': ['SUBSECRETARI', 'OFICIALIA MAYOR'],
                'J': ['JEFE DE UNIDAD', 'JEFATURA DE UNIDAD'],
                'K': ['DIRECTOR GENERAL', 'DIRECCION GENERAL'],
                'L': ['DIRECTOR GENERAL ADJUNT', 'DIRECCION GENERAL ADJUNT'],
                'M': ['DIRECTOR DE AREA', 'DIRECCION DE AREA', 'COORDINADOR', 'COORDINACION'],
                'N': ['SUBDIRECTOR', 'SUBDIRECCION'],
                'O': ['JEFE DE DEPARTAMENTO', 'JEFATURA DE DEPARTAMENTO'],
                'P': ['ENLACE']
            }

            for level, keywords in level_keywords.items():
                if any(keyword in denominacion_upper for keyword in keywords):
                    return level

        # Si no se puede determinar, retornar UNKNOWN
        return 'UNKNOWN'

def robust_openai_call(prompt: str,
                      max_tokens: int = 800,
                      model: str = "openai/gpt-4o",
                      temperature: float = 0.1,
                      context: APFContext = None) -> Dict[str, Any]:
    """
    Llamada robusta a OpenAI con manejo mejorado y logging.
    ADAPTADO PARA V5: Usa OpenAIProvider en lugar de litellm directamente.

    Mantiene la firma original para compatibilidad con validadores v4.
    """
    # Importar OpenAIProvider de v5
    try:
        import sys
        import os
        # Agregar path raíz del proyecto para imports
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        if project_root not in sys.path:
            sys.path.insert(0, project_root)

        from src.providers.openai_provider import OpenAIProvider
        from src.interfaces.llm_provider import LLMRequest, LLMProviderError
    except ImportError as e:
        error_msg = f"No se pudo importar OpenAIProvider de v5: {str(e)}"
        if context:
            context.add_error(error_msg)
        return {"status": "error", "error": error_msg}

    if context:
        context.start_step("openai_call")

    try:
        # Extraer API key del contexto si existe
        api_key = None
        if context and hasattr(context, 'data'):
            api_key = context.data.get('openai_api_key') or context.data.get('api_key')

        # Crear provider de v5
        provider = OpenAIProvider(
            api_key=api_key,
            default_model=model,
            enable_logging=LOGGING_CONFIG.get("log_openai_calls", True)
        )

        # Crear request de v5
        request = LLMRequest(
            prompt=prompt,
            model=model,
            max_tokens=max_tokens,
            temperature=temperature
        )

        if LOGGING_CONFIG.get("log_openai_calls", True):
            print(f"[OpenAI] Llamada iniciada - Model: {model}, Max tokens: {max_tokens}")

        start_time = time.time()

        # Llamar a provider usando complete_json para obtener dict directamente
        try:
            result = provider.complete_json(request)
            duration = time.time() - start_time

            if context:
                context.complete_step("openai_call", f"JSON parseado exitosamente en {duration:.2f}s")

            if LOGGING_CONFIG.get("log_openai_calls", True):
                print(f"[OpenAI] Respuesta recibida en {duration:.2f}s")

            return {
                "status": "success",
                "data": result,
                "metadata": {
                    "model": model,
                    "duration": duration
                }
            }

        except LLMProviderError as e:
            # Si complete_json falla, intentar con complete y parsing manual
            if LOGGING_CONFIG.get("log_openai_calls", True):
                print(f"[OpenAI] complete_json falló, intentando complete normal...")

            response = provider.complete(request)
            duration = time.time() - start_time
            content = response.content

            if not content:
                error_msg = "OpenAI devolvió respuesta vacía"
                if context:
                    context.fail_step("openai_call", error_msg)
                return {"status": "error", "error": error_msg}

            # Intentar parsear como JSON manualmente
            content_cleaned = content.strip()

            # Limpiar markdown wrapper
            if content_cleaned.startswith('```json'):
                content_cleaned = content_cleaned[7:]
                if content_cleaned.endswith('```'):
                    content_cleaned = content_cleaned[:-3]
                content_cleaned = content_cleaned.strip()
            elif content_cleaned.startswith('```'):
                lines = content_cleaned.split('\n')
                if len(lines) > 2 and lines[-1].strip() == '```':
                    content_cleaned = '\n'.join(lines[1:-1])

            try:
                result = json.loads(content_cleaned)
                if context:
                    context.complete_step("openai_call", f"JSON parseado manualmente en {duration:.2f}s")
                return {
                    "status": "success",
                    "data": result,
                    "metadata": {
                        "model": model,
                        "duration": duration
                    }
                }
            except json.JSONDecodeError:
                # Fallback: buscar JSON con regex
                json_patterns = [
                    r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}',
                    r'\[[^\[\]]*(?:\[[^\[\]]*\][^\[\]]*)*\]'
                ]

                for pattern in json_patterns:
                    matches = re.findall(pattern, content_cleaned, re.DOTALL)
                    for match in matches:
                        try:
                            result = json.loads(match)
                            if context:
                                context.complete_step("openai_call", "JSON extraído con regex")
                            return {"status": "success", "data": result}
                        except:
                            continue

                error_msg = f"No se pudo parsear JSON: {str(e)}"
                if context:
                    context.fail_step("openai_call", error_msg)
                return {
                    "status": "partial",
                    "raw_content": content_cleaned,
                    "error": error_msg
                }

    except Exception as e:
        error_msg = f"Error en llamada OpenAI: {str(e)}"
        if context:
            context.fail_step("openai_call", error_msg)
        return {"status": "error", "error": error_msg}

# ==========================================
# PROCESAMIENTO DE DOCUMENTOS UNIFICADO
# ==========================================

def read_pdf_document(file_path: str, context: APFContext = None) -> Dict[str, Any]:
    """Lee un archivo PDF y extrae su contenido"""
    if not PYMUPDF_AVAILABLE:
        error_msg = "PyMuPDF no disponible para procesamiento PDF"
        if context:
            context.add_error(error_msg)
        return {"status": "error", "error": error_msg}
    
    if context:
        context.start_step("pdf_processing")
    
    try:
        doc = fitz.open(file_path)
        content = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            content += page.get_text()
        
        metadata = {
            "total_pages": len(doc),
            "title": doc.metadata.get("title", ""),
            "author": doc.metadata.get("author", ""),
            "file_size": os.path.getsize(file_path)
        }
        
        doc.close()
        
        if context:
            context.complete_step("pdf_processing", f"Extraídas {len(doc)} páginas")
        
        return {
            "content": clean_text_for_processing(content),
            "metadata": metadata,
            "file_type": "pdf",
            "status": "success"
        }
        
    except Exception as e:
        error_msg = f"Error procesando PDF: {str(e)}"
        if context:
            context.fail_step("pdf_processing", error_msg)
        return {"status": "error", "error": error_msg}

def read_docx_document(file_path: str, context: APFContext = None) -> Dict[str, Any]:
    """Lee un archivo DOCX y extrae su contenido"""
    if not DOCX_AVAILABLE:
        error_msg = "python-docx no disponible para procesamiento DOCX"
        if context:
            context.add_error(error_msg)
        return {"status": "error", "error": error_msg}
    
    if context:
        context.start_step("docx_processing")
    
    try:
        doc = Document(file_path)
        content = ""
        
        # Extraer párrafos
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        
        # Extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += cell.text + "\t"
                content += "\n"
        
        metadata = {
            "total_paragraphs": len(doc.paragraphs),
            "total_tables": len(doc.tables),
            "file_size": os.path.getsize(file_path)
        }
        
        if context:
            context.complete_step("docx_processing", f"Extraídos {len(doc.paragraphs)} párrafos")
        
        return {
            "content": clean_text_for_processing(content),
            "metadata": metadata,
            "file_type": "docx",
            "status": "success"
        }
        
    except Exception as e:
        error_msg = f"Error procesando DOCX: {str(e)}"
        if context:
            context.fail_step("docx_processing", error_msg)
        return {"status": "error", "error": error_msg}

def read_text_document(file_path: str, context: APFContext = None) -> Dict[str, Any]:
    """Lee un archivo de texto plano"""
    if context:
        context.start_step("text_processing")
    
    try:
        # Intentar diferentes encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        content = ""
        encoding_used = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                encoding_used = encoding
                break
            except UnicodeDecodeError:
                continue
        
        if not content:
            # Fallback: leer como binario
            with open(file_path, 'rb') as f:
                raw_content = f.read()
                content = raw_content.decode('utf-8', errors='replace')
                encoding_used = 'utf-8-replace'
        
        metadata = {
            "character_count": len(content),
            "line_count": content.count('\n'),
            "file_size": os.path.getsize(file_path),
            "encoding_used": encoding_used
        }
        
        if context:
            context.complete_step("text_processing", f"Archivo leído con encoding {encoding_used}")
        
        return {
            "content": clean_text_for_processing(content),
            "metadata": metadata,
            "file_type": "txt",
            "status": "success"
        }
        
    except Exception as e:
        error_msg = f"Error procesando archivo de texto: {str(e)}"
        if context:
            context.fail_step("text_processing", error_msg)
        return {"status": "error", "error": error_msg}

def detect_file_type_and_read(file_path: str, context: APFContext = None) -> Dict[str, Any]:
    """
    Detecta automáticamente el tipo de archivo y lee su contenido.
    Versión unificada con soporte para contexto.
    """
    if not Path(file_path).exists():
        error_msg = f"Archivo no encontrado: {file_path}"
        if context:
            context.add_error(error_msg)
        return {"status": "error", "error": error_msg}
    
    file_extension = Path(file_path).suffix.lower()
    
    if LOGGING_CONFIG["log_file_processing"]:
        print(f"[FILE] Procesando: {Path(file_path).name} ({file_extension})")
    
    if file_extension == '.pdf':
        return read_pdf_document(file_path, context)
    elif file_extension in ['.docx', '.doc']:
        return read_docx_document(file_path, context)
    elif file_extension in ['.txt', '.text']:
        return read_text_document(file_path, context)
    else:
        # Intentar como texto plano por defecto
        if context:
            context.add_warning(f"Tipo de archivo desconocido: {file_extension}. Procesando como texto.")
        return read_text_document(file_path, context)

# ==========================================
# UTILIDADES DE VALIDACIÓN
# ==========================================

def validate_json_structure(data: Dict[str, Any], required_fields: List[str]) -> Dict[str, Any]:
    """
    Valida estructura JSON con campos requeridos - CORREGIDO DIVISIÓN POR CERO
    """
    missing_fields = []
    present_fields = []
    
    for field_path in required_fields:
        value = safe_get_nested(data, field_path)
        
        if value is not None and value != "" and value != []:
            present_fields.append(field_path)
        else:
            missing_fields.append(field_path)
    
    # CORRECCIÓN: Proteger contra división por cero
    if len(required_fields) == 0:
        completion_percentage = 100.0  # Si no hay campos requeridos, consideramos 100% completo
    else:
        completion_percentage = (len(present_fields) / len(required_fields)) * 100
    
    return {
        "is_valid": len(missing_fields) == 0,
        "completion_percentage": round(completion_percentage, 2),
        "present_fields": present_fields,
        "missing_fields": missing_fields,
        "total_required": len(required_fields),
        "total_present": len(present_fields)
    }

def safe_get_nested(data: Dict[str, Any], path: str, default: Any = None) -> Any:
    """
    Obtiene valor anidado de diccionario de forma segura.
    
    Args:
        data: Diccionario fuente
        path: Ruta con puntos (ej: "identificacion.nivel.codigo")
        default: Valor por defecto si no se encuentra
    """
    try:
        keys = path.split('.')
        current = data
        
        for key in keys:
            if isinstance(current, dict) and key in current:
                current = current[key]
            else:
                return default
        
        return current
    except:
        return default

def calculate_processing_stats(data: Dict[str, Any]) -> Dict[str, Any]:
    """Calcula estadísticas básicas de datos procesados - CORREGIDO DIVISIÓN POR CERO"""
    stats = {
        "timestamp": datetime.now().isoformat(),
        "has_identification": bool(safe_get_nested(data, "identificacion_puesto.denominacion_puesto")),
        "has_objective": bool(safe_get_nested(data, "objetivo_general.descripcion_completa")),
        "functions_count": len(safe_get_nested(data, "funciones", [])),
        "data_quality": "unknown"
    }
    
    # CORRECCIÓN: Manejar caso donde functions_count puede ser 0
    quality_score = 0
    if stats["has_identification"]: quality_score += 1
    if stats["has_objective"]: quality_score += 1
    if stats["functions_count"] > 0: quality_score += 1
    
    # CORRECCIÓN: Verificar que tenemos elementos para evaluar
    if quality_score >= 3:
        stats["data_quality"] = "high"
    elif quality_score >= 2:
        stats["data_quality"] = "medium"
    elif quality_score >= 1:
        stats["data_quality"] = "low"
    else:
        stats["data_quality"] = "none"  # Nuevo estado para documentos sin datos
    
    return stats

# ==========================================
# UTILIDADES DE SISTEMA
# ==========================================

def check_system_dependencies() -> Dict[str, Any]:
    """Verifica dependencias del sistema"""
    return {
        "litellm": LITELLM_AVAILABLE,
        "pymupdf": PYMUPDF_AVAILABLE,
        "docx": DOCX_AVAILABLE,
        "openai_key": bool(os.getenv("OPENAI_API_KEY")),
        "system_ready": all([
            LITELLM_AVAILABLE,
            os.getenv("OPENAI_API_KEY")
        ])
    }

def format_timestamp(dt: datetime = None) -> str:
    """Genera timestamp formateado"""
    if dt is None:
        dt = datetime.now()
    return dt.strftime("%Y-%m-%d %H:%M:%S")

def print_system_status():
    """Imprime estado del sistema"""
    deps = check_system_dependencies()
    
    print("=" * 50)
    print("SISTEMA APF v4.0 - ESTADO DE DEPENDENCIAS")
    print("=" * 50)
    
    for dep, status in deps.items():
        status_icon = "✅" if status else "❌"
        print(f"   {status_icon} {dep}")
    
    if deps["system_ready"]:
        print("\n✅ Sistema listo para operación")
    else:
        print("\n⚠️  Sistema requiere configuración adicional")
        if not deps["openai_key"]:
            print("   - Configura OPENAI_API_KEY")
        if not deps["litellm"]:
            print("   - Instala: pip install litellm")
    
    print("=" * 50)

# ==========================================
# TESTING Y DEBUGGING
# ==========================================

def create_test_context(test_data: Dict[str, Any] = None) -> APFContext:
    """Crea contexto de prueba con datos simulados"""
    context = APFContext("TEST_CONTEXT")
    
    if test_data:
        for key, value in test_data.items():
            context.set_data(key, value, "test_agent")
    
    # Simular algunos pasos completados
    context.start_step("test_step_1", "test_agent")
    context.complete_step("test_step_1", "Test completado exitosamente")
    
    return context

def run_system_tests() -> bool:
    """Ejecuta tests básicos del sistema"""
    print("🧪 EJECUTANDO TESTS DEL SISTEMA APF v4.0")
    print("=" * 50)
    
    try:
        # Test 1: Dependencias
        print("Test 1: Verificando dependencias...")
        deps = check_system_dependencies()
        if not deps["system_ready"]:
            print("❌ Test 1 falló - dependencias faltantes")
            return False
        print("✅ Test 1 pasó")
        
        # Test 2: Contexto
        print("Test 2: Probando contexto...")
        context = create_test_context({"test_key": "test_value"})
        if context.get_data("test_key") != "test_value":
            print("❌ Test 2 falló - problema con contexto")
            return False
        print("✅ Test 2 pasó")
        
        # Test 3: Procesamiento de texto
        print("Test 3: Probando limpieza de texto...")
        dirty_text = "Test\x00text\nwith\x1cproblems"
        clean = clean_text_for_processing(dirty_text)
        if '\x00' in clean or '\x1c' in clean:
            print("❌ Test 3 falló - limpieza de texto")
            return False
        print("✅ Test 3 pasó")
        
        # Test 4: Extracción de nivel jerárquico
        print("Test 4: Probando extracción jerárquica...")
        level = extract_hierarchical_level(nivel_salarial="M33")
        if level != "M":
            print("❌ Test 4 falló - extracción jerárquica")
            return False
        print("✅ Test 4 pasó")
        
        print("\n✅ TODOS LOS TESTS PASARON")
        return True
        
    except Exception as e:
        print(f"❌ Error en tests: {str(e)}")
        return False

if __name__ == "__main__":
    print_system_status()
    print("\n")
    run_system_tests()